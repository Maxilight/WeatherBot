'''
import re
s = 'asd/asd/asd/asd/asd/asd/asd/as'
result = re.match('asd', s)
print(result[0])
s = 'asd/asd/asd/asd/asd/asd/asd/as'
result = re.search('/as', s)
print(result[0])
s = 'asd/asd/asd/asd/asd/asd/asd/as'
result = re.findall('/as', s)
print(result)
s = 'asd/asd/asd/asd/asd/asd/asd/as'
result = re.split('/', s)
print(result)
s = 'asd/asd/asd/asd/asd/asd/asd/as'
result = re.sub('as', 'qq', s)
print(result)
stringa = '43---234423 %KKK1asdasd hr +375295467789 +375441567780 +375331233123'
result = re.findall(r'\d\d\d\d\d\d\d\d', stringa)
print(result)
stringa = '43---234423 %KKK1asdasd hr +375295467789 +375441567780 +375331233123'
result = re.findall(r'\D\D\D\D\D', stringa)
print(result)
stringa = '43---234423 %KKK1asdasd hr +375295467789 +375441567780 +375331233123'
result = re.findall(r'\s', stringa)
print(result)
'''
import docx
from docx import Document
from docx.shared import Inches
document = Document()
document.add_heading('Заголовок, ')
p = document.add_paragraph('Параграф ')
p.add_run('Жирный ').bold = True
p.add_run('Курсивный ').italic = True
document.save('document')

